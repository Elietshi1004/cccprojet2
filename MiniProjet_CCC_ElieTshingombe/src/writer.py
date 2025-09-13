#!/usr/bin/env python3
"""
Module de génération des documents de sortie
Gère la création des documents Word, CSV et Excel formatés

Ce module contient toutes les fonctions nécessaires pour :
1. Générer des documents Word avec structure hiérarchique
2. Créer des exports CSV et Excel normalisés
3. Appliquer le formatage professionnel (couleurs, tableaux)
4. Gérer la signature automatique et la traçabilité

Auteur: ElieTshingombe
Date: 2025
Projet: Mini-projet CCC - Automatisation du traitement des RAW DATA
"""

import pandas as pd
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from datetime import datetime
from utils import file_hash

def export_csv(data, output_path):
    """
    Exporte les données au format CSV
    
    Args:
        data (list): Liste des dictionnaires de mesures
        output_path (str): Chemin de sortie du fichier CSV
    """
    df = pd.DataFrame.from_records(data)  # ✅ gère colonnes manquantes
    df.to_csv(output_path, index=False)

def export_xlsx(data, output_path):
    """
    Exporte les données au format Excel (.xlsx)
    
    Args:
        data (list): Liste des dictionnaires de mesures
        output_path (str): Chemin de sortie du fichier Excel
    """
    df = pd.DataFrame.from_records(data)
    df.to_excel(output_path, index=False)

def export_word(test_params, data, summary, global_verdict,
                output_path, candidate_name, raw_path):

    doc = Document()

    # 🔸 Cartouche en-tête avec paramètres de test
    doc.add_heading("Test Parameters", level=1)
    keys = ["Sample ID", "RBW", "Antenne", "Orientation DUT", "Operator", "Date/Heure", "Conclusion"]
    table = doc.add_table(rows=0, cols=2)
    table.style = 'Table Grid'
    for k in keys:
        row = table.add_row().cells
        row[0].text = k
        row[1].text = str(test_params.get(k, "—"))

    # 🔸 Tableaux normalisés par section
    if data:
        # Grouper par Sample ID et Configuration
        grouped_data = group_by_sample_and_config(data, test_params)
        
        for sample_id, configs in grouped_data.items():
            doc.add_heading(f"Sample n°{sample_id}", level=1)
            
            for config_name, measurements in configs.items():
                doc.add_heading(f"Configuration {config_name}", level=2)
                
                # Créer le tableau avec les colonnes attendues selon questions.txt
                table = doc.add_table(rows=1, cols=9)
                table.style = 'Table Grid'
                
                # En-têtes selon les spécifications
                headers = [
                    "Section", "Frequency (MHz)", "SR", "Polarization", 
                    "Correction (dB)", "Mesure (dBµV/m)", "Limite (dBµV/m)", 
                    "Marge (dB)", "Verdict"
                ]
                
                hdr = table.rows[0].cells
                for i, header in enumerate(headers):
                    hdr[i].text = header
                    run = hdr[i].paragraphs[0].runs[0]
                    run.bold = True
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Texte noir
                
                # Ajouter les données
                for row in measurements:
                    cells = table.add_row().cells
                    
                    # Section (détecteur type)
                    cells[0].text = str(row.get("Detector type", "-"))
                    
                    # Frequency (MHz) - formatage selon spécifications
                    freq = row.get("Frequency (MHz)", "")
                    if isinstance(freq, (int, float)):
                        if freq < 10:
                            cells[1].text = f"{freq:.5f}"
                        else:
                            cells[1].text = f"{freq:.3f}"
                    else:
                        cells[1].text = str(freq)
                    
                    # SR (pas spécifié dans les données, mettre par défaut)
                    cells[2].text = "-"
                    
                    # Polarization
                    cells[3].text = str(row.get("Polarization", "Vertical"))
                    
                    # Correction (dB) - pas dans les données, mettre par défaut
                    cells[4].text = "0.00"
                    
                    # Mesure (dBµV/m)
                    cells[5].text = str(row.get("Mesure (dBµV/m)", "-"))
                    
                    # Limite (dBµV/m)
                    cells[6].text = str(row.get("Limite (dBµV/m)", "-"))
                    
                    # Marge (dB) - formatage 2 décimales
                    margin = row.get("Margin (dB)", "-")
                    if isinstance(margin, (int, float)):
                        cells[7].text = f"{margin:.2f}"
                    else:
                        cells[7].text = str(margin)
                    
                    # Verdict avec couleur
                    verdict_txt = str(row.get("Conformity", "-"))
                    run = cells[8].paragraphs[0].add_run(verdict_txt)
                    if verdict_txt.upper() == "OK":
                        run.font.color.rgb = RGBColor(0, 128, 0)  # vert
                    elif verdict_txt.upper() == "NOK":
                        run.font.color.rgb = RGBColor(200, 0, 0)  # rouge
                        run.bold = True
    else:
        doc.add_paragraph("Aucune mesure disponible.")


    # 🔸 Synthèse finale
    doc.add_heading("Synthèse finale", level=1)
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Section"
    hdr[1].text = "Nb lignes"
    hdr[2].text = "Nb FAIL"
    hdr[3].text = "Verdict"
    for row in summary:
        cells = table.add_row().cells
        cells[0].text = str(row["Section"])
        cells[1].text = str(row["NbLines"])
        cells[2].text = str(row["NbFAIL"])
        verdict_txt = str(row["Verdict"])
        run = cells[3].paragraphs[0].add_run(verdict_txt)
        if verdict_txt == "OK":
            run.font.color.rgb = RGBColor(0, 128, 0)
        else:
            run.font.color.rgb = RGBColor(200, 0, 0)
            run.bold = True

    doc.add_paragraph(f"Verdict global : {global_verdict}")

    # 🔸 Pied de page signature
    h = file_hash(raw_path)
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    footer_text = f"{candidate_name} | {now} | {h}"
    section = doc.sections[0]
    footer = section.footer
    if len(footer.paragraphs) == 0:
        p = footer.add_paragraph()
    else:
        p = footer.paragraphs[0]
    p.text = footer_text

    doc.save(output_path)


def export_word_multiple_samples(all_samples_data, all_processed_data, all_summaries,
                                output_path, candidate_name, raw_path):
    """
    Export Word pour plusieurs Sample ID
    """
    doc = Document()

    # 🔸 Titre principal
    doc.add_heading("1.1 Test results", level=1)
    doc.add_paragraph("The result table mentions only the worst cases. For the details see complete tables in the measurements and curves.")

    # 🔸 Pour chaque Sample ID
    for sample_id, sample_data in all_samples_data.items():
        test_params = sample_data['test_params']
        configurations = sample_data['configurations']
        sample_processed_data = all_processed_data.get(sample_id, {})
        sample_summaries = all_summaries.get(sample_id, {})
        
        # Titre du Sample
        doc.add_heading(f"Sample n°{sample_id}", level=1)
        
        # Pour chaque configuration de ce Sample ID
        for config in configurations:
            config_name = config['config_name']
            processed = sample_processed_data.get(config_name, [])
            summary, global_verdict = sample_summaries.get(config_name, ([], "NOK"))
            
            print(f"Configuration {config_name}: {len(processed)} mesures")
            
            # Titre de la configuration
            doc.add_heading(f"Configuration {config_name}", level=2)
            
            # Créer le tableau avec les colonnes attendues
            table = doc.add_table(rows=1, cols=9)
            table.style = 'Table Grid'
            
            # En-têtes selon les spécifications
            headers = [
                "Section", "Frequency (MHz)", "SR", "Polarization", 
                "Correction (dB)", "Mesure (dBµV/m)", "Limite (dBµV/m)", 
                "Marge (dB)", "Verdict"
            ]
            
            hdr = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr[i].text = header
                run = hdr[i].paragraphs[0].runs[0]
                run.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)  # Texte noir
            
            # Ajouter les données
            for row in processed:
                if isinstance(row, dict):  # Vérifier que c'est un dictionnaire
                    cells = table.add_row().cells
                    
                    # Section (détecteur type)
                    cells[0].text = str(row.get("Detector type", "-"))
                    
                    # Frequency (MHz) - formatage selon spécifications
                    freq = row.get("Frequency (MHz)", "")
                    if isinstance(freq, (int, float)):
                        if freq < 10:
                            cells[1].text = f"{freq:.5f}"
                        else:
                            cells[1].text = f"{freq:.3f}"
                    else:
                        cells[1].text = str(freq)
                    
                    # SR (pas spécifié dans les données, mettre par défaut)
                    cells[2].text = "-"
                    
                    # Polarization
                    cells[3].text = str(row.get("Polarization", "Vertical"))
                    
                    # Correction (dB) - pas dans les données, mettre par défaut
                    cells[4].text = "0.00"
                    
                    # Mesure (dBµV/m)
                    cells[5].text = str(row.get("Mesure (dBµV/m)", "-"))
                    
                    # Limite (dBµV/m)
                    cells[6].text = str(row.get("Limite (dBµV/m)", "-"))
                    
                    # Marge (dB) - formatage 2 décimales
                    margin = row.get("Margin (dB)", "-")
                    if isinstance(margin, (int, float)):
                        cells[7].text = f"{margin:.2f}"
                    else:
                        cells[7].text = str(margin)
                    
                    # Verdict avec couleur
                    verdict_txt = str(row.get("Conformity", "-"))
                    run = cells[8].paragraphs[0].add_run(verdict_txt)
                    if verdict_txt.upper() == "OK":
                        run.font.color.rgb = RGBColor(0, 128, 0)  # vert
                    elif verdict_txt.upper() == "NOK":
                        run.font.color.rgb = RGBColor(200, 0, 0)  # rouge
                        run.bold = True
                else:
                    print(f"Erreur: row n'est pas un dictionnaire: {type(row)} - {row}")
            
            # Si pas de données pour cette configuration, afficher un message
            if len(processed) == 0:
                doc.add_paragraph("Aucune donnée disponible pour cette configuration.")

    # 🔸 Pied de page signature
    h = file_hash(raw_path)
    now = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    footer_text = f"{candidate_name} | {now} | {h}"
    section = doc.sections[0]
    footer = section.footer
    if len(footer.paragraphs) == 0:
        p = footer.add_paragraph()
    else:
        p = footer.paragraphs[0]
    p.text = footer_text

    doc.save(output_path)


def group_by_sample_and_config(data, test_params):
    """
    Groupe les mesures par Sample ID et Configuration
    """
    grouped = {}
    
    # Récupérer le Sample ID
    sample_id = test_params.get("Sample ID", "Unknown")
    
    # Déterminer la configuration (RBW + Mode)
    rbw = test_params.get("RBW", "9kHz")
    mode = test_params.get("Operating mode", "Mode 3")
    config_name = f"({mode}) RBW {rbw}"
    
    if sample_id not in grouped:
        grouped[sample_id] = {}
    
    if config_name not in grouped[sample_id]:
        grouped[sample_id][config_name] = []
    
    # S'assurer que data est une liste de dictionnaires
    if isinstance(data, list):
        grouped[sample_id][config_name].extend(data)
    else:
        print(f"Erreur: data n'est pas une liste: {type(data)}")
    
    return grouped


def group_measurements_by_position(data):
    """
    Groupe les mesures par position d'antenne et polarisation
    """
    grouped = {}
    
    for row in data:
        # Extraire position d'antenne et polarisation
        antenna_pos = row.get("Antenna Position", "1 (X)")
        polarization = row.get("Polarization", "Vertical")
        
        key = (antenna_pos, polarization)
        if key not in grouped:
            grouped[key] = []
        grouped[key].append(row)
    
    return grouped
