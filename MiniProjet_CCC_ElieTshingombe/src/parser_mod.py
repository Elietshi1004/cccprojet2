#!/usr/bin/env python3
"""
Module d'extraction des données depuis les documents Word RAW
Gère l'extraction et la normalisation des données CEM/EMI

Ce module contient toutes les fonctions nécessaires pour :
1. Extraire les Sample ID et configurations depuis les documents Word
2. Parser les tableaux de mesures (CISPR.AVG, Peak, Q-Peak)
3. Associer les données de mesures à leurs configurations
4. Normaliser et nettoyer les données extraites

Auteur: ElieTshingombe
Date: 2025
Projet: Mini-projet CCC - Automatisation du traitement des RAW DATA
"""

from docx import Document
import re

from utils import clean_decimal, normalize_unit, normalize_header

def extract_data(filepath):
    """
    Fonction principale d'extraction des données depuis un fichier Word RAW
    
    Cette fonction orchestre l'ensemble du processus d'extraction :
    1. Détecte toutes les configurations dans le document
    2. Groupe les configurations par Sample ID
    3. Extrait les paramètres de test pour chaque Sample ID
    4. Extrait les mesures pour chaque configuration
    5. Retourne une structure hiérarchique complète
    
    Args:
        filepath (str): Chemin vers le fichier Word RAW (.docx)
    
    Returns:
        dict: Structure hiérarchique des données extraites
            {
                sample_id: {
                    'test_params': dict,           # Paramètres de test
                    'configurations': list,        # Liste des configurations
                    'config_measurements': dict    # Mesures par configuration
                }
            }
    """
    # Charger le document Word
    doc = Document(filepath)
    
    # ========================================================================
    # ÉTAPE 1: DÉTECTION DE TOUTES LES CONFIGURATIONS
    # ========================================================================
    
    # Extraire toutes les configurations du fichier (Sample ID + nom de configuration)
    all_configurations = extract_all_configurations(doc)
    print(f"Configurations trouvées : {len(all_configurations)}")
    for config in all_configurations:
        print(f"  - {config['sample_id']} : {config['config_name']}")
    
    # ========================================================================
    # ÉTAPE 2: GROUPEMENT PAR SAMPLE ID
    # ========================================================================
    
    # Grouper les configurations par Sample ID
    sample_groups = {}
    for config in all_configurations:
        sample_id = config['sample_id']
        if sample_id not in sample_groups:
            sample_groups[sample_id] = []
        sample_groups[sample_id].append(config)
    
    # ========================================================================
    # ÉTAPE 3: EXTRACTION DES DONNÉES PAR SAMPLE ID
    # ========================================================================
    
    all_samples_data = {}
    
    # Traiter chaque Sample ID séparément
    for sample_id, configs in sample_groups.items():
        print(f"\nTraitement du Sample ID : {sample_id} ({len(configs)} configurations)")
        
        # Extraire les paramètres de test pour ce Sample ID
        test_params = extract_test_params_for_sample(doc, sample_id)
        
        # Extraire les mesures pour chaque configuration de ce Sample ID
        config_measurements = {}
        for config in configs:
            config_name = config['config_name']
            print(f"  Configuration : {config_name}")
            
            # Extraire les mesures pour cette configuration spécifique
            measurements = extract_measurements_for_configuration(doc, sample_id, config_name)
            config_measurements[config_name] = measurements
        
        # Stocker toutes les données de ce Sample ID
        all_samples_data[sample_id] = {
            'test_params': test_params,           # Paramètres de test (RBW, antenne, etc.)
            'configurations': configs,            # Liste des configurations disponibles
            'config_measurements': config_measurements  # Mesures par configuration
        }
    
    return all_samples_data


def extract_all_sample_ids(doc):
    """
    Extrait tous les Sample ID du document
    """
    sample_ids = set()
    
    # Chercher dans les tableaux "name test:"
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) >= 2:
                key = row.cells[0].text.strip().lower()
                value = row.cells[1].text.strip()
                
                if "name test:" in key and value:
                    # Format: CRE2-2025-TP002-02_ER_In front of harness RBW 9kHz
                    sample_id = value.split("_")[0]
                    
                    # Filtrer les Sample ID valides (format: CRE2-2025-TP002-XX)
                    if re.match(r'^CRE2-\d{4}-TP\d{3}-\d{2}$', sample_id):
                        sample_ids.add(sample_id)
                    elif "CRE2-2025-TP002" in sample_id:
                        # Extraire le vrai Sample ID si format différent
                        match = re.search(r'CRE2-\d{4}-TP\d{3}-\d{2}', sample_id)
                        if match:
                            sample_ids.add(match.group())
    
    return list(sample_ids)


def extract_all_configurations(doc):
    """
    Extrait toutes les configurations (Sample ID + Configuration) du document
    """
    configurations = []
    
    # Chercher dans les tableaux "name test:"
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) >= 2:
                key = row.cells[0].text.strip().lower()
                value = row.cells[1].text.strip()
                
                if "name test:" in key and value:
                    # Format: CRE2-2025-TP002-02_ER_In front of harness RBW 9kHz
                    parts = value.split("_")
                    if len(parts) >= 2:
                        sample_id = parts[0]
                        config_name = "_".join(parts[1:])  # Tout après le premier _
                        
                        # Filtrer les Sample ID valides
                        if re.match(r'^CRE2-\d{4}-TP\d{3}-\d{2}$', sample_id):
                            configurations.append({
                                'sample_id': sample_id,
                                'config_name': config_name,
                                'full_name': value
                            })
    
    return configurations


def extract_test_params_for_sample(doc, sample_id):
    """
    Extrait les paramètres de test pour un Sample ID spécifique
    """
    test_params = {"Sample ID": sample_id}
    
    # Chercher les paramètres dans les tableaux
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) >= 2:
                key = row.cells[0].text.strip().lower()
                value = row.cells[1].text.strip()
                
                # Extraire Project
                if "project:" in key and value and value != "Project:":
                    test_params["Project"] = value
                
                # Extraire RBW depuis les tableaux de configuration
                elif "rbw" in key and value and value != "RBW:":
                    test_params["RBW"] = value
                
                # Extraire Operator
                elif "operator:" in key and value and value != "Operator:":
                    test_params["Operator"] = value
                
                # Extraire Test Configuration
                elif "test configuration:" in key and value and value != "Test Configuration:":
                    test_params["Test Configuration"] = value
                
                # Extraire Operating mode
                elif "operating mode:" in key and value and value != "Operating mode:":
                    test_params["Operating mode"] = value
                
                # Extraire Conclusion
                elif "conclusion" in key and value and value != "Conclusion:":
                    test_params["Conclusion"] = value
    
    return test_params


def extract_measurements_for_sample(doc, sample_id):
    """
    Extrait les mesures pour un Sample ID spécifique depuis les tableaux Word
    """
    measurements = []
    
    # Chercher les tableaux de mesures qui correspondent à ce Sample ID
    for table_idx, table in enumerate(doc.tables):
        # Vérifier si ce tableau contient des données pour ce Sample ID
        table_text = " ".join([cell.text for row in table.rows for cell in row.cells])
        
        print(f"Tableau {table_idx}: {len(table.rows)} lignes")
        print(f"  Contient {sample_id}? {sample_id in table_text}")
        
        if sample_id in table_text:
            print(f"  *** Tableau {table_idx} contient {sample_id} ***")
            # Extraire les mesures de ce tableau
            table_measurements = extract_measurements_from_table(table, sample_id)
            if isinstance(table_measurements, list):
                measurements.extend(table_measurements)
                print(f"  Ajouté {len(table_measurements)} mesures")
            else:
                print(f"Erreur: table_measurements n'est pas une liste: {type(table_measurements)}")
        else:
            # Vérifier si c'est un tableau de mesures (avec Frequency, Limit, etc.)
            if len(table.rows) > 1:
                headers = [cell.text.strip() for cell in table.rows[0].cells]
                if any("Frequency" in h or "Limit" in h or "Margin" in h for h in headers):
                    print(f"  Tableau {table_idx} semble être un tableau de mesures (pas lié à {sample_id})")
                    # Essayer d'extraire quand même
                    table_measurements = extract_measurements_from_table(table, sample_id)
                    if isinstance(table_measurements, list) and len(table_measurements) > 0:
                        measurements.extend(table_measurements)
                        print(f"  Ajouté {len(table_measurements)} mesures (tableau générique)")
    
    print(f"Mesures extraites pour {sample_id}: {len(measurements)}")
    return measurements


def extract_measurements_for_configuration(doc, sample_id, config_name):
    """
    Extrait les mesures pour une configuration spécifique
    
    LOGIQUE D'EXTRACTION :
    1. Trouve le tableau "Test parameters" contenant le nom de la configuration
    2. Remonte dans les tableaux précédents pour trouver les tableaux de mesures
    3. Extrait les données de tous les tableaux de mesures trouvés
    4. Associe les mesures à la configuration correspondante
    
    Cette approche est nécessaire car dans les documents Word RAW :
    - Les tableaux de mesures (CISPR.AVG, Peak, Q-Peak) précèdent les "Test parameters"
    - Chaque configuration a ses propres tableaux de mesures
    - Il faut identifier la bonne association configuration ↔ mesures
    
    Args:
        doc (Document): Document Word chargé
        sample_id (str): Identifiant du Sample (ex: "CRE2-2025-TP002-02")
        config_name (str): Nom de la configuration (ex: "ER_In front of harness RBW 9kHz")
    
    Returns:
        list: Liste des mesures extraites pour cette configuration
    """
    measurements = []
    
    # ========================================================================
    # ÉTAPE 1: RECHERCHE DU TABLEAU "TEST PARAMETERS"
    # ========================================================================
    
    # Parcourir tous les tableaux pour trouver celui contenant cette configuration
    for table_idx, table in enumerate(doc.tables):
        table_text = " ".join([cell.text for row in table.rows for cell in row.cells])
        
        # Vérifier si c'est un tableau "Test parameters" avec cette configuration
        if "name test:" in table_text.lower() and config_name in table_text:
            print(f"  *** Tableau {table_idx} est Test parameters pour {config_name} ***")
            
            # ========================================================================
            # ÉTAPE 2: EXTRACTION DES TABLEAUX DE MESURES PRÉCÉDENTS
            # ========================================================================
            
            # Chercher TOUS les tableaux de mesures qui précèdent ce tableau
            # On remonte depuis le tableau "Test parameters" vers le début
            for check_idx in range(table_idx - 1, -1, -1):  # De table_idx-1 vers 0
                check_table = doc.tables[check_idx]
                check_headers = [cell.text.strip() for cell in check_table.rows[0].cells]
                
                # Vérifier si c'est un tableau de mesures (contient Frequency, Limit, Margin)
                if any("Frequency" in h or "Limit" in h or "Margin" in h for h in check_headers):
                    print(f"  *** Tableau {check_idx} est un tableau de mesures ***")
                    
                    # Extraire les mesures de ce tableau
                    table_measurements = extract_measurements_from_table(check_table, sample_id)
                    if isinstance(table_measurements, list) and len(table_measurements) > 0:
                        measurements.extend(table_measurements)
                        print(f"  Ajouté {len(table_measurements)} mesures du tableau {check_idx}")
                else:
                    # Si ce n'est pas un tableau de mesures, on s'arrête
                    # (on a atteint un autre type de tableau)
                    print(f"  *** Tableau {check_idx} n'est pas un tableau de mesures, arrêt ***")
                    break
    
    print(f"Mesures extraites pour {sample_id} - {config_name}: {len(measurements)}")
    return measurements


def extract_measurements_from_table(table, sample_id):
    """
    Extrait les mesures d'un tableau spécifique pour un Sample ID
    """
    measurements = []
    
    if len(table.rows) < 2:
        return measurements
    
    headers = [normalize_header(normalize_unit(c.text.strip())) for c in table.rows[0].cells]

    if any("Frequency" in h or "Limit" in h or "Margin" in h for h in headers):
        for row in table.rows[1:]:
            values = [c.text.strip() for c in row.cells]
            if not any(values):
                continue

            row_dict = {}
            for i, h in enumerate(headers):
                val = values[i] if i < len(values) else ""
                if re.search(r"\d", val):
                    try:
                        row_dict[h] = clean_decimal(val)
                    except:
                        row_dict[h] = val
                else:
                    row_dict[h] = val
            
            # Ajouter les colonnes manquantes avec des valeurs par défaut
            row_dict["Polarization"] = "Vertical"   # Valeur par défaut
            row_dict["Comment"] = "-"
            row_dict["Sample ID"] = sample_id  # Ajouter le Sample ID à chaque mesure
            
            # Normaliser les noms de colonnes pour correspondre au format attendu
            for h in headers:
                print(f"En-tête trouvé : {h}")
                if "Frequency" in h and "MHz" not in h:
                    row_dict["Frequency (MHz)"] = row_dict.pop(h, "")
                elif "Detector" in h:
                    row_dict["Detector type"] = row_dict.pop(h, "")
                
                # Mapper les colonnes selon les spécifications
                if "cispr" in h.lower() and "avg" in h.lower() and "lim" not in h.lower():
                    row_dict["Mesure (dBµV/m)"] = row_dict.pop(h, "")
                    row_dict["Detector type"] = "CISPR.AVG"
                elif "q-peak" in h.lower() and "lim" not in h.lower():
                    row_dict["Mesure (dBµV/m)"] = row_dict.pop(h, "")
                    row_dict["Detector type"] = "Q-Peak"
                elif "peak" in h.lower() and "lim" not in h.lower() and "q-peak" not in h.lower():
                    row_dict["Mesure (dBµV/m)"] = row_dict.pop(h, "")
                    row_dict["Detector type"] = "Peak"
                elif "lim" in h.lower() and "avg" in h.lower():
                    row_dict["Limite (dBµV/m)"] = row_dict.pop(h, "")
                elif "lim" in h.lower() and "q-peak" in h.lower():
                    row_dict["Limite (dBµV/m)"] = row_dict.pop(h, "")
                elif "lim" in h.lower() and "peak" in h.lower():
                    row_dict["Limite (dBµV/m)"] = row_dict.pop(h, "")
                elif "cispr" in h.lower() and "lim" in h.lower() and "avg" in h.lower():
                    row_dict["Margin (dB)"] = row_dict.pop(h, "")
                elif "peak" in h.lower() and "lim" in h.lower() and "q-peak" in h.lower():
                    row_dict["Margin (dB)"] = row_dict.pop(h, "")
                elif "q-peak" in h.lower() and "lim" in h.lower():
                    row_dict["Margin (dB)"] = row_dict.pop(h, "")
                elif "peak" in h.lower() and "lim" in h.lower() and "peak" in h.lower():
                    row_dict["Margin (dB)"] = row_dict.pop(h, "")
            
            measurements.append(row_dict)

    return measurements


def extract_measurements(doc):
    """
    Fonction legacy - à supprimer après migration
    """
    measurements = []
    return measurements
