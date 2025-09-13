#!/usr/bin/env python3
"""
Module d'extraction des données depuis les documents Word RAW
Gère l'extraction et la normalisation des données CEM/EMI

Ce module contient toutes les fonctions nécessaires pour :
1. Extraire les Sample ID et configurations depuis les documents Word
2. Parser les tableaux de mesures (CISPR.AVG, Peak, Q-Peak)
3. Associer les données de mesures à leurs configurations
4. Normaliser et nettoyer les données extraites

Auteur: ElieTshingombe
Date: 2025
Projet: Mini-projet CCC - Automatisation du traitement des RAW DATA
"""

from docx import Document
import re

from utils import clean_decimal, normalize_unit, normalize_header

def extract_data(filepath):
    """
    Fonction principale d'extraction des données depuis un fichier Word RAW
    
    Cette fonction orchestre l'ensemble du processus d'extraction :
    1. Détecte toutes les configurations dans le document
    2. Groupe les configurations par Sample ID
    3. Extrait les paramètres de test pour chaque Sample ID
    4. Extrait les mesures pour chaque configuration
    5. Retourne une structure hiérarchique complète
    
    Args:
        filepath (str): Chemin vers le fichier Word RAW (.docx)
    
    Returns:
        dict: Structure hiérarchique des données extraites
            {
                sample_id: {
                    'configurations': list,        # Liste des configurations
                    'config_measurements': dict,   # Mesures par configuration
                    'config_test_params': dict     # Paramètres de test par configuration
                }
            }
    """
    # Charger le document Word
    doc = Document(filepath)

    # ========================================================================
    # ÉTAPE 1: DÉTECTION DE TOUTES LES CONFIGURATIONS
    # ========================================================================
    
    # Extraire toutes les configurations du fichier (Sample ID + nom de configuration)
    all_configurations = extract_all_configurations(doc)
    print(f"Configurations trouvées : {len(all_configurations)}")
    for config in all_configurations:
        print(f"  - {config['sample_id']} : {config['config_name']}")
    
    # ========================================================================
    # ÉTAPE 2: GROUPEMENT PAR SAMPLE ID
    # ========================================================================
    
    # Grouper les configurations par Sample ID
    sample_groups = {}
    for config in all_configurations:
        sample_id = config['sample_id']
        if sample_id not in sample_groups:
            sample_groups[sample_id] = []
        sample_groups[sample_id].append(config)
    
    # ========================================================================
    # ÉTAPE 3: EXTRACTION DES DONNÉES PAR SAMPLE ID
    # ========================================================================
    
    all_samples_data = {}
    
    # Traiter chaque Sample ID séparément
    for sample_id, configs in sample_groups.items():
        print(f"\nTraitement du Sample ID : {sample_id} ({len(configs)} configurations)")
        
        # Extraire les mesures et paramètres pour chaque configuration
        config_measurements = {}
        config_test_params = {}
        
        for config in configs:
            config_name = config['config_name']
            print(f"  Configuration : {config_name}")
            
            # Extraire les paramètres de test pour cette configuration
            test_params = extract_test_params_for_configuration(doc, sample_id, config_name)
            config_test_params[config_name] = test_params
            
            # Extraire les mesures pour cette configuration
            measurements = extract_measurements_for_configuration(doc, sample_id, config_name)
            config_measurements[config_name] = measurements
        
        # Stocker toutes les données de ce Sample ID
        all_samples_data[sample_id] = {
            'configurations': configs,                    # Liste des configurations disponibles
            'config_measurements': config_measurements,   # Mesures par configuration
            'config_test_params': config_test_params      # Paramètres de test par configuration
        }
    
    return all_samples_data


def extract_all_sample_ids(doc):
    """
    Extrait tous les Sample ID du document
    """
    sample_ids = set()
    
    # Chercher dans les tableaux "name test:"
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) >= 2:
                key = row.cells[0].text.strip().lower()
                value = row.cells[1].text.strip()
                
                if "name test:" in key and value:
                    # Format: CRE2-2025-TP002-02_ER_In front of harness RBW 9kHz
                    sample_id = value.split("_")[0]
                    
                    # Filtrer les Sample ID valides (format: CRE2-2025-TP002-XX)
                    if re.match(r'^CRE2-\d{4}-TP\d{3}-\d{2}$', sample_id):
                        sample_ids.add(sample_id)
                    elif "CRE2-2025-TP002" in sample_id:
                        # Extraire le vrai Sample ID si format différent
                        match = re.search(r'CRE2-\d{4}-TP\d{3}-\d{2}', sample_id)
                        if match:
                            sample_ids.add(match.group())
    
    return list(sample_ids)


def extract_all_configurations(doc):
    """
    Extrait toutes les configurations (Sample ID + Configuration) du document
    """
    configurations = []
    
    # Chercher dans les tableaux "name test:"
    for table in doc.tables:
        for row in table.rows:
            if len(row.cells) >= 2:
                key = row.cells[0].text.strip().lower()
                value = row.cells[1].text.strip()
                
                if "name test:" in key and value:
                    # Format: CRE2-2025-TP002-02_ER_In front of harness RBW 9kHz
                    parts = value.split("_")
                    if len(parts) >= 2:
                        sample_id = parts[0]
                        config_name = "_".join(parts[1:])  # Tout après le premier _
                        
                        # Filtrer les Sample ID valides
                        if re.match(r'^CRE2-\d{4}-TP\d{3}-\d{2}$', sample_id):
                            configurations.append({
                                'sample_id': sample_id,
                                'config_name': config_name,
                                'full_name': value
                            })
    
    return configurations


def extract_test_params_for_configuration(doc, sample_id, config_name, debug=False):
    """
    Extrait les paramètres de test pour une configuration spécifique.
    Robuste :
     - gère clé/valeur sur 2 colonnes (cell[0]=clé, cell[1]=valeur)
     - gère clé: valeur dans une seule cellule (cell[0] contient "Operator: John")
     - nettoie espaces non imprimables (NBSP, zero-width, tab)
     - debug optionnel pour afficher le contenu des cellules
    Args:
        doc (Document): objet python-docx
        sample_id (str): id du sample (ex: "CRE2-2025-TP002-02")
        config_name (str): nom de la configuration (ex: "ER_In front of harness RBW 9kHz")
        debug (bool): si True, affiche beaucoup d'informations utiles pour debug
    Retour:
        dict test_params
    """
    def normalize_text(s):
        if s is None:
            return ""
        # remplacer NBSP et caractères invisibles, normaliser espaces et retours
        s = s.replace("\xa0", " ")
        s = s.replace("\u200b", "")   # zero width space
        s = s.replace("\t", " ")
        s = s.replace("\r", " ")
        s = s.replace("\n", " ")
        # collapse spaces
        s = " ".join(s.split())
        return s.strip()

    def normalize_key(k):
        if k is None:
            return ""
        return normalize_text(k).lower().replace(":", "").strip()

    # résultat initial
    test_params = {
        "Sample ID": sample_id,
        "Configuration": config_name
    }

    if debug:
        print("DEBUG: recherche des tableaux contenant des clés de paramètres (sample/project/operator) ...")

    # trouver tables candidates (celles qui contiennent au moins une des clés cherchées)
    candidate_tables = []
    keywords = ["sample", "project", "operator", "test configuration", "operating mode", "conclusion", "rbw", "span", "reference level"]
    for ti, table in enumerate(doc.tables):
        # concat des textes de toutes les cellules (normalisé)
        table_text = " ".join(normalize_text(cell.text) for row in table.rows for cell in row.cells).lower()
        if any(kw in table_text for kw in keywords):
            candidate_tables.append((ti, table, table_text))
            if debug:
                print(f"  candidate table index={ti} (preview): {table_text[:200]}")

    if not candidate_tables and debug:
        print("DEBUG: aucune table candidate trouvée contenant les mots-clés. On retournera le sample_id seulement.")

    # Parcourir les tables candidates et extraire lignes clé/valeur
    found_any = False
    for ti, table, _ in candidate_tables:
        if debug:
            print(f"\nDEBUG: inspection table index={ti} - {len(table.rows)} lignes")
        # parcourir les lignes de la table
        for ri, row in enumerate(table.rows):
            # récupérer textes de chaque cellule normalisés
            cells = [normalize_text(c.text) for c in row.cells]
            if debug:
                print(f"  ROW {ri}: {cells!r}")

            # ignorer lignes vides
            if not any(cells):
                if debug:
                    print("    -> ligne vide, skip")
                continue

            # cas 1 : valeur présente dans cell[1]
            key_text = ""
            value_text = ""
            if len(cells) >= 2 and cells[1].strip():
                key_text = cells[0]
                value_text = cells[1]
            else:
                # cas 2 : tout dans cell[0] comme "Operator: NDN/WD, 17/02/2025..."
                cell0 = cells[0]
                if ":" in cell0:
                    parts = cell0.split(":", 1)
                    # si la "clé" est très courte (ex: 'Operator') on prend la partie après ':'
                    key_text = parts[0]
                    value_text = parts[1]
                else:
                    # cas 3 : essayer de trouver une valeur dans d'autres colonnes (cell[2], cell[3], ...)
                    key_text = cells[0]
                    value_text = ""
                    for j in range(1, len(cells)):
                        if cells[j].strip():
                            value_text = cells[j]
                            break

            # nettoyer
            key_norm = normalize_key(key_text)
            value_norm = normalize_text(value_text)

            if debug:
                print(f"    parsed -> key='{key_norm}', value='{value_norm}'")

            # si valeur vide, on ignore cette ligne
            if not value_norm:
                if debug:
                    print("    -> pas de valeur trouvée pour cette ligne, skip")
                continue

            found_any = True

            # mapping des clés (cherchez des sous-chaînes, ce qui rend la détection tolérante)
            if "sample" in key_norm:
                test_params["Sample ID"] = value_norm
            elif "project" in key_norm:
                test_params["Project"] = value_norm
            elif "operator" in key_norm:
                test_params["Operator"] = value_norm
            elif "test configuration" in key_norm or "test cfg" in key_norm:
                test_params["Test Configuration"] = value_norm
            elif "operating mode" in key_norm or key_norm.startswith("mode"):
                # gérer "Mode 3, Conclusion: comply" dans la même valeur
                low = value_norm.lower()
                if "conclusion" in low:
                    parts = value_norm.split("conclusion", 1)
                    test_params["Operating mode"] = parts[0].replace(":", "").strip().rstrip(",")
                    # extraire la conclusion si présente après 'conclusion'
                    rest = parts[1].replace(":", "").strip()
                    if rest:
                        test_params["Conclusion"] = rest
                else:
                    test_params["Operating mode"] = value_norm
            elif "conclusion" in key_norm:
                test_params["Conclusion"] = value_norm
            elif "rbw" == key_norm or "rbw" in key_norm:
                test_params["RBW"] = value_norm
            elif "span" in key_norm:
                test_params["Span"] = value_norm
            elif "reference" in key_norm and "level" in key_norm:
                test_params["Reference level"] = value_norm
            else:
                # si aucune clé reconnue, on peut stocker sur un dictionnaire 'OtherParams'
                test_params.setdefault("OtherParams", {})[key_norm] = value_norm

        # si on a déjà trouvé des paramètres utiles dans cette table, on peut retourner
        if found_any:
            if debug:
                print(f"DEBUG: paramètres extraits depuis la table index={ti}: {test_params}")
            return test_params

    # Si aucune table candidate n'a donné de valeur, faire une passe globale (fallback) : chercher lignes "Key: Value" dans TOUTES les cellules
    if debug:
        print("DEBUG: fallback - scan global de toutes les cellules pour 'Key: Value'")

    for ti, table in enumerate(doc.tables):
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                txt = normalize_text(cell.text)
                if ":" in txt:
                    parts = txt.split(":", 1)
                    k = normalize_key(parts[0])
                    v = normalize_text(parts[1])
                    if not v:
                        continue
                    if "project" in k and "Project" not in test_params:
                        test_params["Project"] = v
                    elif "operator" in k and "Operator" not in test_params:
                        test_params["Operator"] = v
                    elif "test configuration" in k and "Test Configuration" not in test_params:
                        test_params["Test Configuration"] = v
                    elif "operating mode" in k and "Operating mode" not in test_params:
                        test_params["Operating mode"] = v
                    elif "conclusion" in k and "Conclusion" not in test_params:
                        test_params["Conclusion"] = v

    if debug:
        print(f"DEBUG: résultat final (fallback): {test_params}")

    return test_params


# Fonction supprimée - remplacée par extract_measurements_for_configuration


def extract_measurements_for_configuration(doc, sample_id, config_name):
    """
    Extrait les mesures pour une configuration spécifique
    
    NOUVELLE LOGIQUE :
    1. Trouve le tableau "Test parameters" contenant cette configuration
    2. Remonte dans les tableaux précédents pour trouver les tableaux de mesures
    3. Extrait les données de tous les tableaux de mesures trouvés
    
    Args:
        doc (Document): Document Word chargé
        sample_id (str): Identifiant du Sample (ex: "CRE2-2025-TP002-02")
        config_name (str): Nom de la configuration (ex: "ER_In front of harness RBW 9kHz")
    
    Returns:
        list: Liste des mesures extraites pour cette configuration
    """
    measurements = []
    
    print(f"  Recherche des mesures pour {config_name}")
    
    # ========================================================================
    # ÉTAPE 1: TROUVER LE TABLEAU "TEST PARAMETERS" POUR CETTE CONFIGURATION
    # ========================================================================
    
    target_table_idx = None
    for table_idx, table in enumerate(doc.tables):
        table_text = " ".join([cell.text for row in table.rows for cell in row.cells])
        if "name test:" in table_text.lower() and config_name in table_text:
            target_table_idx = table_idx
            print(f"  *** Tableau Test parameters trouvé à l'index {table_idx} ***")
            break
    
    if target_table_idx is None:
        print(f"  ❌ Aucun tableau Test parameters trouvé pour {config_name}")
        return measurements
    
    # ========================================================================
    # ÉTAPE 2: EXTRAIRE LES TABLEAUX DE MESURES PRÉCÉDENTS
    # ========================================================================
    
    # Remonter depuis le tableau "Test parameters" vers le début
    # Mais s'arrêter au premier tableau de paramètres trouvé pour éviter les doublons
    for check_idx in range(target_table_idx - 1, -1, -1):
        check_table = doc.tables[check_idx]
        check_headers = [cell.text.strip() for cell in check_table.rows[0].cells]
        
        print(f"  Vérification du tableau {check_idx}: {check_headers}")
        
        # Vérifier si c'est un tableau de paramètres (Sample, Project, etc.)
        is_params_table = any(keyword in h for h in check_headers for keyword in 
            ["Sample:", "Project:", "Operator:", "Test Configuration:", "Operating mode:"])
        
        if is_params_table:
            print(f"  *** Tableau {check_idx} est un tableau de paramètres, arrêt de la remontée ***")
            break
        
        # Vérifier si c'est un tableau de mesures
        is_measurement_table = any(keyword in h for h in check_headers for keyword in 
            ["Frequency", "CISPR", "Peak", "Q-Peak", "Lim.Avg", "Lim.Q-Peak", "Lim.Peak"])
        
        if is_measurement_table:
            print(f"  *** Tableau {check_idx} est un tableau de mesures ***")
            
            # Extraire les mesures de ce tableau
            table_measurements = extract_measurements_from_table(check_table, sample_id)
            if isinstance(table_measurements, list) and len(table_measurements) > 0:
                # Éviter les doublons en vérifiant si la mesure existe déjà
                for measure in table_measurements:
                    # Créer une clé unique basée sur la fréquence et la mesure
                    freq = measure.get("Frequency (MHz)", "")
                    mes = measure.get("Mesure (dBµV/m)", "")
                    key = f"{freq}_{mes}"
                    
                    # Vérifier si cette mesure existe déjà
                    existing_keys = [f"{m.get('Frequency (MHz)', '')}_{m.get('Mesure (dBµV/m)', '')}" for m in measurements]
                    if key not in existing_keys:
                        measurements.append(measure)
                    else:
                        print(f"    Mesure dupliquée ignorée: {freq} MHz, {mes} dBµV/m")
                
                print(f"  Ajouté {len(table_measurements)} mesures du tableau {check_idx}")
            else:
                print(f"  Aucune mesure extraite du tableau {check_idx}")
        else:
            # Si ce n'est pas un tableau de mesures, on s'arrête
            print(f"  *** Tableau {check_idx} n'est pas un tableau de mesures, arrêt ***")
            break
    
    print(f"  ✅ Mesures extraites pour {config_name}: {len(measurements)}")
    return measurements


def extract_measurements_from_table(table, sample_id):
    """
    Extrait les mesures d'un tableau spécifique pour un Sample ID.
    Corrigé pour être plus robuste :
    - Fix indentation
    - Détection souple des headers
    - Mapping basé sur mots-clés
    """
    measurements = []

    # Vérifier que le tableau a au moins une ligne d'en-têtes + 1 ligne de données
    if len(table.rows) < 2:
        print(f"    Tableau trop petit: {len(table.rows)} lignes")
        return measurements

    # Extraire les en-têtes normalisés
    headers = [normalize_header(normalize_unit(c.text.strip())) for c in table.rows[0].cells]
    print(f"    En-têtes du tableau: {headers}")

    # Vérifier que c'est bien un tableau de mesures
    if not any(kw in h.lower() for h in headers for kw in ["frequency", "cispr", "peak", "q-peak", "lim", "margin"]):
        print("    *** Tableau ne contient pas de mesures ***")
        return measurements

    print(f"    *** Tableau contient {len(table.rows)-1} lignes de données ***")

    # Parcourir les lignes de données
    for row_idx, row in enumerate(table.rows[1:]):
        values = [c.text.strip() for c in row.cells]

        if not any(values):
            print(f"      Ligne {row_idx+1}: vide, ignorée")
            continue

        row_dict = {
            "Sample ID": sample_id,
            "Comment": "-",
            "Section": "-"
        }

        for i, h in enumerate(headers):
            val = values[i] if i < len(values) else ""
            h_low = h.lower()

            print(f"    Mapping colonne {i}: '{h}' -> '{h_low}' -> Valeur: '{val}'")

            if "frequency" in h_low:
                print(f"      -> Frequency: {val}")
                row_dict["Frequency (MHz)"] = clean_decimal(val)
            elif h_low == "sr" or h_low.strip() == "sr":
                print(f"      -> SR: {val}")
                row_dict["S R"] = clean_decimal(val)
            elif "cispr" in h_low and "avg" in h_low and "lim" not in h_low:
                print(f"      -> CISPR.AVG: {val}")
                row_dict["Mesure (dBµV/m)"] = clean_decimal(val)
                row_dict["Detector type"] = "CISPR.AVG"
                row_dict["Section"] = "CISPR.AVG"
            elif "q-peak" in h_low and "lim" not in h_low:
                print(f"      -> Q-Peak: {val}")
                row_dict["Mesure (dBµV/m)"] = clean_decimal(val)
                row_dict["Detector type"] = "Q-Peak"
                row_dict["Section"] = "Q-Peak"
            elif "peak" in h_low and "lim" not in h_low and "q-peak" not in h_low and "-" not in h_low:
                print(f"      -> Peak: {val}")
                row_dict["Mesure (dBµV/m)"] = clean_decimal(val)
                row_dict["Detector type"] = "Peak"
                row_dict["Section"] = "Peak"
            elif ("lim" in h_low or "limit" in h_low) and ("avg" in h_low or "peak" in h_low or "q-peak" in h_low):
                print(f"      -> Limite: {val}")
                row_dict["Limite (dBµV/m)"] = clean_decimal(val)
            elif "margin" in h_low or ("-" in h_low and ("peak" in h_low or "avg" in h_low or "q-peak" in h_low)):
                print(f"      -> Margin: {val}")
                row_dict["Margin (dB)"] = clean_decimal(val)
            elif "pol" in h_low:  # Polarisation
                print(f"      -> Polarisation: {val}")
                row_dict["Polarization"] = val
            elif "corr" in h_low:  # Correction
                print(f"      -> Correction: {val}")
                row_dict["Correction (dB)"] = clean_decimal(val)
            else:
                print(f"      -> Autre colonne: {h} = {val}")
                row_dict[h] = val

        # Ajouter la ligne extraite
        measurements.append(row_dict)
        print(f"      ✅ Ligne {row_idx+1} extraite: {row_dict}")

    print(f"    Total mesures extraites: {len(measurements)}")
    return measurements


def extract_measurements(doc):
    """
    Fonction legacy - à supprimer après migration
    """
    measurements = []
    return measurements
