#!/usr/bin/env python3
"""
Script de debug complet pour diagnostiquer les problèmes d'extraction
"""

from parser_mod import extract_data
from docx import Document

def debug_complete():
    """Debug complet de l'extraction des données"""
    print("=== DEBUG COMPLET DE L'EXTRACTION ===")
    
    # 1. Vérifier la structure du document
    print("\n1. ANALYSE DU DOCUMENT WORD")
    doc = Document("raw/raw01.docx")
    print(f"Nombre de tableaux dans le document : {len(doc.tables)}")
    
    # Analyser chaque tableau
    for i, table in enumerate(doc.tables):
        table_text = " ".join([cell.text for row in table.rows for cell in row.cells])
        print(f"\nTableau {i}:")
        print(f"  Lignes : {len(table.rows)}")
        print(f"  Colonnes : {len(table.rows[0].cells) if table.rows else 0}")
        print(f"  Contenu (200 premiers caractères) : {table_text[:200]}...")
        
        # Vérifier les en-têtes
        if table.rows:
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            print(f"  En-têtes : {headers}")
    
    # 2. Tester l'extraction des configurations
    print("\n2. EXTRACTION DES CONFIGURATIONS")
    from parser_mod import extract_all_configurations
    configurations = extract_all_configurations(doc)
    print(f"Configurations trouvées : {len(configurations)}")
    for config in configurations:
        print(f"  - {config['sample_id']} : {config['config_name']}")
    
    # 3. Tester l'extraction des données
    print("\n3. EXTRACTION DES DONNÉES")
    data = extract_data("raw/raw01.docx")
    
    for sample_id, sample_data in data.items():
        print(f"\nSample ID: {sample_id}")
        config_measurements = sample_data['config_measurements']
        config_test_params = sample_data['config_test_params']
        
        print(f"  Configurations : {len(config_measurements)}")
        for config_name, measurements in config_measurements.items():
            print(f"    {config_name}: {len(measurements)} mesures")
            
        print(f"  Paramètres de test : {len(config_test_params)}")
        for config_name, params in config_test_params.items():
            print(f"    {config_name}: {params}")

if __name__ == "__main__":
    debug_complete()
